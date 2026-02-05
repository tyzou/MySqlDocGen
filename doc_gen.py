import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pymysql
import datetime
import ctypes
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DBDocGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("MySQL 数据库文档生成器")
        self.root.geometry("680x750")

        # Fix for Windows Taskbar Icon
        try:
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("mycompany.dbdocgen.app.1.0")
        except Exception:
            pass

        # Data
        self.conn = None
        self.all_tables = []

        # Style configuration
        self.set_icon()
        self.setup_styles()
        self.setup_ui()

    def set_icon(self):
        # Embed a simple database icon (32x32 PNG Base64)
        # This is a cleaner, more standard icon string
        icon_base64 = (
            "iVBORw0KGgoAAAANSUhEUgAAACAAAAAgCAYAAABzenr0AAAABGdBTUEAALGPC/xhBQAA"
            "ACBjSFJNAAB6JgAAgIQAAPoAAACA6AAAdTAAAOpgAAA6mAAAF3CculE8AAAABmJLR0QA"
            "/wD/AP+gvaeTAAAAB3RJTUUH5gIFBQ4yJ0qgDAAABwJJREFUWMOll1tsFNcZx39nZnb2"
            "Mmvv2uvuGqwNvoDdxCEQJ6A0MY1J05S2D61UValSVaV9aCq1DwVIpKpSqyhCqtRKbdRI"
            "eQil5SGKg0NwaVwMxGAPODa+7/ra3bXXuzO7M9OHXQ/eXXtNVPxJo5lz5j//7/vO+c43"
            "C//nIf5fA/09d/F25Q1eLW9g8y6iFEXhR64f+3j5Y57/bB3F93/A+u/9kC3uFjY0b8S0LAzL"
            "wrIsDMNAMAQ0TcPX14euaR7wH979M6889wI/e/4FfrbhR/zghz/G/l+D+7tu07C5Hk3X0TQN"
            "QRDQNA1N0xAEAUmSkCQJWZZRFAVFUVAUBUmSkCQJSZK4fOkyR48eQ5Zl3v3gA9R4nJ9s+hE/"
            "2/Qj/rD+B/zwB/W4/qtA74Y63n/3PVRVRVEUVFVFVVUUReF2y7bt27dPCSE+2rRpkxBCyLIs"
            "y7Isy/KxY8eEEELu7u4W27dvF4Ig8M67f6SqqorXf7ae537yAj97/gXW/7AO1wOB9W+8ga7r"
            "aJqGpmnouo6u6xiG4S1+w4YNiqIokUik+Nprr0mKokhCiL/X1dUpqqpKIyMj0uTkpBQMBqWi"
            "KFIwGJSmp6elYDAo/e1vfwuHw2FFURS2bNnC5cuXCQaD/GH9D3j5tVd55bnn+cmGjdwLBPbe"
            "dw9d19F1HV3X0XUd0zS9xXfu3CkpivLhV199JXRdlzRNk3RdD4dCoVA4HA6Hw+FwOBwO67qu"
            "S4qikMlkpMHBQWn//v3S/v37Q4FAQAiCgGVZ7L7vXjrv2oZhGBiGgWEYXheYpsmpU6ckz/OC"
            "b7/9tiDLsixJUsjS9WA0Gg1Go9FgNBwOGoaBqqocO3ZM2rNnj7R3796QJEmSJEmSJEmSJEmS"
            "JO26665wR0dHuL+/Pzg1NSUYFkW+27bx+1fuQ3TNDFN01t8fHycf/zjH4LnecGvfvWrYCaTk"
            "WRZliQpYOl6MBqJBqPhSNAwDBRF4fDhw9Kbb74p7dmzJyRJkiRJkiRJkiRJ0q5du8IdHR3h"
            "/v7+4OTkpGBZFvv37eP2L9+GaZqYpolpmliWheM4/OMf/5A8z/MFg8GgLMuyJEkBS9eD0Ug0"
            "GA1Hg4ZhoCgKhw8flnbs2CHt6eoKSZIkSZIkSZIkSZKkHTt2hDs6OsL9/f3BiYkJwbIs9u/d"
            "y+fvuB3TNDFN01t8fHycf/zjH4LnecGvfvWrYCaTkWRZliQpYOl6MBqJBqPhSNAwDBRF4fDhw"
            "9Kbb74p7dmzJyRJkiRJkiRJkiRJ0u7du8IdHR3h/v7+4Pj4uGBZFvv37uW2227DNE1M08S2b"
            "ezp6SmcO3dO8Dwv+MUvfhHMZDKSLMuyJAUswwgyDCOo67qgKAqHDx+WduzYIe3Zs0eSJEmSJ"
            "EmSJEmSJEm7d+8Od3R0hPv7+4Pj4+OCZVnsvW0vX/j87V6X2LaNZVm4rss//vEPwfM8X5ZlW"
            "ZIk3zTMoK7rgqIoHD58WNqxY4e0Z88eSZIkSZIkSZIkSZK0e/fucEdHR3h/v7+4NjYmWJbF3"
            "Xv20HrrFzBN07tA4Louf/vbnwXP83xZlmVJknzTMIOGYQiKonD48GFpx44d0p49eyRJkiRJk"
            "iRJkiRJ0u7du8MdHR3h/v7+4NjYmGBZFvvu2Mvtt9+OaZrYto3jOLiuy9///nfB8zxfQJIkX"
            "5Zl3zTMoGEYgqIoHDp0SNq5c6e0d+9eSZIkSZIkSZIkSZK0e/fucEdHR3h/v7+4NjYmWJbFv"
            "jv2ctttX8A0Te/K8b/2gWEY3qqPjIyE+/r6guPj44Jt29x9517u/PxtmKaJZVnYto3jOLiu6"
            "xV/4oknguPj44Jt29x151f44he+gGma3pXjvwR4V47/sf/j/wH1f9420rtjEQAAAABJRU5Er"
            "kJggg=="
        )
        try:
            img_data = tk.PhotoImage(data=icon_base64)
            self.root.iconphoto(True, img_data)
        except Exception:
            pass

    def setup_styles(self):
        style = ttk.Style()
        # Windows native look is usually default, but let's ensure fonts are nice
        default_font = ("Microsoft YaHei UI", 10)
        
        style.configure(".", font=default_font)
        style.configure("TLabel", font=default_font, padding=2)
        style.configure("TButton", font=default_font, padding=4)
        style.configure("TEntry", font=default_font, padding=2)
        style.configure("TLabelframe", font=("Microsoft YaHei UI", 10, "bold"))
        style.configure("TLabelframe.Label", font=("Microsoft YaHei UI", 10, "bold"), foreground="#333333")

    def setup_ui(self):
        # Main Container with padding
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.pack(fill="both", expand=True)

        # --- 1. Connection Frame ---
        conn_frame = ttk.LabelFrame(main_frame, text=" 数据库连接配置 ", padding=(15, 10))
        conn_frame.pack(fill="x", pady=(0, 15))

        # Grid layout for connection inputs
        # Column 0 & 2: Labels, Column 1 & 3: Entries
        conn_frame.columnconfigure(1, weight=1)
        conn_frame.columnconfigure(3, weight=1)

        ttk.Label(conn_frame, text="主机地址:").grid(row=0, column=0, sticky="w", padx=(0, 5), pady=5)
        self.host_var = tk.StringVar(value="localhost")
        ttk.Entry(conn_frame, textvariable=self.host_var).grid(row=0, column=1, sticky="ew", padx=(0, 15), pady=5)

        ttk.Label(conn_frame, text="端口号:").grid(row=0, column=2, sticky="w", padx=(0, 5), pady=5)
        self.port_var = tk.StringVar(value="3306")
        ttk.Entry(conn_frame, textvariable=self.port_var).grid(row=0, column=3, sticky="ew", pady=5)

        ttk.Label(conn_frame, text="用户名:").grid(row=1, column=0, sticky="w", padx=(0, 5), pady=5)
        self.user_var = tk.StringVar(value="root")
        ttk.Entry(conn_frame, textvariable=self.user_var).grid(row=1, column=1, sticky="ew", padx=(0, 15), pady=5)

        ttk.Label(conn_frame, text="密码:").grid(row=1, column=2, sticky="w", padx=(0, 5), pady=5)
        self.pwd_var = tk.StringVar()
        ttk.Entry(conn_frame, textvariable=self.pwd_var, show="*").grid(row=1, column=3, sticky="ew", pady=5)

        ttk.Label(conn_frame, text="数据库名:").grid(row=2, column=0, sticky="w", padx=(0, 5), pady=5)
        self.db_var = tk.StringVar()
        ttk.Entry(conn_frame, textvariable=self.db_var).grid(row=2, column=1, sticky="ew", padx=(0, 15), pady=5)

        # Connect Button (Spanning a bit or placed strategically)
        self.btn_connect = ttk.Button(conn_frame, text="连接数据库", command=self.connect_db)
        self.btn_connect.grid(row=2, column=2, columnspan=2, sticky="ew", pady=5)


        # --- 2. Table Selection Frame ---
        table_frame = ttk.LabelFrame(main_frame, text=" 数据表选择 ", padding=(15, 10))
        table_frame.pack(fill="both", expand=True, pady=(0, 15))

        # Filter Area
        filter_frame = ttk.Frame(table_frame)
        filter_frame.pack(fill="x", pady=(0, 10))
        
        ttk.Label(filter_frame, text="搜索表名:").pack(side="left", padx=(0, 5))
        self.filter_var = tk.StringVar()
        self.filter_var.trace("w", self.filter_tables)
        ttk.Entry(filter_frame, textvariable=self.filter_var).pack(side="left", fill="x", expand=True)

        # Listbox Area
        list_container = ttk.Frame(table_frame)
        list_container.pack(fill="both", expand=True, pady=(0, 10))
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")

        # Listbox font
        self.listbox = tk.Listbox(
            list_container, 
            selectmode="extended", 
            yscrollcommand=scrollbar.set, 
            font=("Consolas", 11),
            activestyle="none",
            borderwidth=1,
            relief="solid"
        )
        self.listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=self.listbox.yview)

        # Action Buttons for Selection
        btn_frame = ttk.Frame(table_frame)
        btn_frame.pack(fill="x")
        ttk.Button(btn_frame, text="全选", command=self.select_all).pack(side="left", padx=(0, 10))
        ttk.Button(btn_frame, text="清空选择", command=self.clear_selection).pack(side="left")
        
        # Counter label
        self.count_label = ttk.Label(btn_frame, text="共 0 张表", foreground="gray")
        self.count_label.pack(side="right")


        # --- 3. Action Frame (Bottom) ---
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill="x")
        
        # Status
        self.status_var = tk.StringVar(value="请连接数据库...")
        status_lbl = ttk.Label(action_frame, textvariable=self.status_var, foreground="#666666", font=("Microsoft YaHei UI", 9))
        status_lbl.pack(side="left", anchor="center")
        
        # Main Action Buttons
        self.btn_gen_word = ttk.Button(action_frame, text="导出 Word", command=self.generate_doc, state="disabled")
        self.btn_gen_word.pack(side="right", padx=2)
        
        self.btn_gen_html = ttk.Button(action_frame, text="导出 HTML", command=self.generate_html, state="disabled")
        self.btn_gen_html.pack(side="right", padx=2)
        
        self.btn_gen_md = ttk.Button(action_frame, text="导出 Markdown", command=self.generate_md, state="disabled")
        self.btn_gen_md.pack(side="right", padx=2)

    def connect_db(self):
        try:
            self.conn = pymysql.connect(
                host=self.host_var.get(),
                port=int(self.port_var.get()),
                user=self.user_var.get(),
                password=self.pwd_var.get(),
                database=self.db_var.get(),
                charset='utf8mb4',
                cursorclass=pymysql.cursors.DictCursor
            )
            self.status_var.set("已连接数据库。")
            self.btn_gen_word.config(state="normal")
            self.btn_gen_html.config(state="normal")
            self.btn_gen_md.config(state="normal")
            self.fetch_tables()
        except Exception as e:
            messagebox.showerror("连接错误", str(e))
            self.status_var.set("连接失败。")
            self.btn_gen_word.config(state="disabled")
            self.btn_gen_html.config(state="disabled")
            self.btn_gen_md.config(state="disabled")

    def fetch_tables(self):
        if not self.conn:
            return
        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SHOW TABLE STATUS")
                tables = cursor.fetchall()
                # Store tuples of (Name, Comment)
                self.all_tables = [(t['Name'], t['Comment']) for t in tables]
                self.update_listbox(self.all_tables)
        except Exception as e:
            messagebox.showerror("错误", f"获取表列表失败: {e}")

    def update_listbox(self, items):
        self.listbox.delete(0, tk.END)
        for name, comment in items:
            display_text = f"{name}"
            if comment:
                display_text += f" ({comment})"
            self.listbox.insert(tk.END, display_text)
        
        # Update count label
        self.count_label.config(text=f"共 {len(items)} 张表")

    def filter_tables(self, *args):
        search_term = self.filter_var.get().lower()
        filtered = [t for t in self.all_tables if search_term in t[0].lower() or search_term in t[1].lower()]
        self.update_listbox(filtered)

    def select_all(self):
        self.listbox.select_set(0, tk.END)

    def clear_selection(self):
        self.listbox.selection_clear(0, tk.END)

    def get_column_info(self, table_name):
        # Fetch detailed column info
        # Using information_schema for better standard access to types and comments
        sql = """
            SELECT 
                COLUMN_NAME, 
                COLUMN_TYPE,
                DATA_TYPE,
                IS_NULLABLE, 
                COLUMN_DEFAULT, 
                COLUMN_COMMENT 
            FROM information_schema.COLUMNS 
            WHERE TABLE_SCHEMA = %s AND TABLE_NAME = %s 
            ORDER BY ORDINAL_POSITION
        """
        with self.conn.cursor() as cursor:
            cursor.execute(sql, (self.db_var.get(), table_name))
            return cursor.fetchall()

    def get_selected_table_data(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("警告", "请至少选择一张表！")
            return None

        selected_data = []
        for idx in selected_indices:
            display_text = self.listbox.get(idx)
            table_name = display_text.split(" (")[0]
            table_comment = next((t[1] for t in self.all_tables if t[0] == table_name), "")
            columns = self.get_column_info(table_name)
            selected_data.append({
                'name': table_name,
                'comment': table_comment,
                'columns': columns
            })
        return selected_data

    def parse_column_type(self, col):
        col_type = col['DATA_TYPE']
        col_len = ""
        full_type = col['COLUMN_TYPE']
        if "(" in full_type and ")" in full_type:
            start = full_type.find("(") + 1
            end = full_type.rfind(")")
            col_len = full_type[start:end]
        return col_type, col_len

    def set_cell_font(self, cell, text, bold=False):
        paragraph = cell.paragraphs[0]
        paragraph.clear() # Clear existing content if any
        run = paragraph.add_run(str(text) if text is not None else "")
        run.bold = bold
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(9)

    def generate_doc(self):
        if not self.conn:
            messagebox.showwarning("警告", "请先连接数据库！")
            return

        table_data = self.get_selected_table_data()
        if not table_data:
            return

        default_filename = f"数据库文档_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        file_path = filedialog.asksaveasfilename(
            initialfile=default_filename,
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="保存数据库文档"
        )
        if not file_path:
            return

        try:
            doc = Document()
            
            # Global Style for Normal text
            style = doc.styles['Normal']
            style.font.name = 'Microsoft YaHei'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            # Title
            title = doc.add_heading(f'数据库设计文档: {self.db_var.get()}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for table_info in table_data:
                t_name = table_info['name']
                t_comment = table_info['comment']
                columns = table_info['columns']

                # Table Heading
                header_text = f"表名: {t_name}"
                if t_comment:
                    header_text += f"   说明: {t_comment}"
                doc.add_heading(header_text, level=2)

                # Create Word Table
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.autofit = False 
                
                # Header Row
                headers = ['字段名', '类型', '长度/定义', '允许空', '默认值', '注释']
                hdr_cells = table.rows[0].cells
                for i, h_text in enumerate(headers):
                    self.set_cell_font(hdr_cells[i], h_text, bold=True)

                # Data Rows
                for col in columns:
                    row_cells = table.add_row().cells
                    col_type, col_len = self.parse_column_type(col)
                    
                    self.set_cell_font(row_cells[0], col['COLUMN_NAME'], bold=True)
                    self.set_cell_font(row_cells[1], col_type) 
                    self.set_cell_font(row_cells[2], col_len)
                    self.set_cell_font(row_cells[3], col['IS_NULLABLE'])
                    self.set_cell_font(row_cells[4], col['COLUMN_DEFAULT'] if col['COLUMN_DEFAULT'] is not None else 'NULL')
                    self.set_cell_font(row_cells[5], col['COLUMN_COMMENT'])

                doc.add_paragraph() # Spacer

            doc.save(file_path)
            messagebox.showinfo("完成", f"文档已生成: {file_path}")
            
        except Exception as e:
            messagebox.showerror("生成失败", f"错误详情: {e}")

    def generate_html(self):
        if not self.conn:
            messagebox.showwarning("警告", "请先连接数据库！")
            return

        table_data = self.get_selected_table_data()
        if not table_data:
            return

        default_filename = f"数据库文档_{datetime.datetime.now().strftime('%Y%m%d')}.html"
        file_path = filedialog.asksaveasfilename(
            initialfile=default_filename,
            defaultextension=".html",
            filetypes=[("HTML Document", "*.html")],
            title="保存 HTML 文档"
        )
        if not file_path:
            return

        try:
            db_name = self.db_var.get()
            html_content = [
                "<!DOCTYPE html>",
                "<html>",
                "<head>",
                f"<title>数据库设计文档: {db_name}</title>",
                "<meta charset='utf-8'>",
                "<style>",
                "body { font-family: 'Microsoft YaHei', sans-serif; margin: 40px; background-color: #f5f5f5; }",
                "h1 { text-align: center; color: #333; }",
                "h2 { margin-top: 40px; color: #444; border-bottom: 2px solid #ddd; padding-bottom: 10px; }",
                "table { width: 100%; border-collapse: collapse; margin-top: 10px; background-color: #fff; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }",
                "th, td { border: 1px solid #ccc; padding: 12px; text-align: left; }",
                "th { background-color: #f8f8f8; font-weight: bold; }",
                "tr:nth-child(even) { background-color: #fafafa; }",
                ".col-name { font-weight: bold; color: #2c3e50; }",
                "</style>",
                "</head>",
                "<body>",
                f"<h1>数据库设计文档: {db_name}</h1>"
            ]

            for table_info in table_data:
                t_name = table_info['name']
                t_comment = table_info['comment']
                columns = table_info['columns']

                header_text = f"表名: {t_name}"
                if t_comment:
                    header_text += f" &nbsp;&nbsp; 说明: {t_comment}"
                
                html_content.append(f"<h2>{header_text}</h2>")
                html_content.append("<table>")
                html_content.append("<thead><tr><th>字段名</th><th>类型</th><th>长度/定义</th><th>允许空</th><th>默认值</th><th>注释</th></tr></thead>")
                html_content.append("<tbody>")

                for col in columns:
                    col_type, col_len = self.parse_column_type(col)
                    default_val = col['COLUMN_DEFAULT'] if col['COLUMN_DEFAULT'] is not None else 'NULL'
                    
                    html_content.append("<tr>")
                    html_content.append(f"<td class='col-name'>{col['COLUMN_NAME']}</td>")
                    html_content.append(f"<td>{col_type}</td>")
                    html_content.append(f"<td>{col_len}</td>")
                    html_content.append(f"<td>{col['IS_NULLABLE']}</td>")
                    html_content.append(f"<td>{default_val}</td>")
                    html_content.append(f"<td>{col['COLUMN_COMMENT']}</td>")
                    html_content.append("</tr>")

                html_content.append("</tbody></table>")

            html_content.append("</body></html>")

            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(html_content))

            messagebox.showinfo("完成", f"HTML 文档已生成: {file_path}")

        except Exception as e:
            messagebox.showerror("生成失败", f"错误详情: {e}")

    def generate_md(self):
        if not self.conn:
            messagebox.showwarning("警告", "请先连接数据库！")
            return

        table_data = self.get_selected_table_data()
        if not table_data:
            return

        default_filename = f"数据库文档_{datetime.datetime.now().strftime('%Y%m%d')}.md"
        file_path = filedialog.asksaveasfilename(
            initialfile=default_filename,
            defaultextension=".md",
            filetypes=[("Markdown Document", "*.md")],
            title="保存 Markdown 文档"
        )
        if not file_path:
            return

        try:
            db_name = self.db_var.get()
            md_content = [
                f"# 数据库设计文档: {db_name}\n"
            ]

            for table_info in table_data:
                t_name = table_info['name']
                t_comment = table_info['comment']
                columns = table_info['columns']

                header_text = f"## 表名: {t_name}"
                if t_comment:
                    header_text += f"   说明: {t_comment}"
                
                md_content.append(header_text)
                md_content.append("\n| 字段名 | 类型 | 长度/定义 | 允许空 | 默认值 | 注释 |")
                md_content.append("| --- | --- | --- | --- | --- | --- |")

                for col in columns:
                    col_type, col_len = self.parse_column_type(col)
                    default_val = col['COLUMN_DEFAULT'] if col['COLUMN_DEFAULT'] is not None else 'NULL'
                    
                    # Escape pipe character in comment or default value if necessary
                    comment = col['COLUMN_COMMENT'].replace("|", "\\|")
                    
                    md_content.append(f"| **{col['COLUMN_NAME']}** | {col_type} | {col_len} | {col['IS_NULLABLE']} | {default_val} | {comment} |")
                
                md_content.append("\n")

            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(md_content))

            messagebox.showinfo("完成", f"Markdown 文档已生成: {file_path}")

        except Exception as e:
            messagebox.showerror("生成失败", f"错误详情: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DBDocGeneratorApp(root)
    root.mainloop()